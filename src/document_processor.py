import os
import json
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Track processed files to avoid duplicates
        self.processed_files = set()
    
    def load_documents(self, directory_path: str) -> List[LangchainDocument]:
        """Load all documents from a directory"""
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.warning(f"Directory {directory_path} does not exist")
            return documents
        
        logger.info(f"Loading documents from: {directory_path}")
        
        # Get all files in directory
        file_paths = []
        for file_path in directory.rglob("*"):
            if file_path.is_file() and self._is_supported_file(file_path):
                file_paths.append(file_path)
        
        logger.info(f"Found {len(file_paths)} supported files")
        
        # Process each file
        for file_path in file_paths:
            try:
                doc = self._load_single_file(file_path)
                if doc:
                    documents.append(doc)
                    logger.info(f"✅ Loaded: {file_path.name}")
                else:
                    logger.warning(f"⚠️ Skipped: {file_path.name} (no content)")
                    
            except Exception as e:
                logger.error(f"❌ Error loading {file_path.name}: {str(e)}")
                continue
        
        logger.info(f"Successfully loaded {len(documents)} documents")
        return documents
    
    def load_single_document(self, file_path: str) -> Optional[LangchainDocument]:
        """Load a single document by file path"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                logger.error(f"File not found: {file_path}")
                return None
            
            if not self._is_supported_file(file_path):
                logger.error(f"Unsupported file type: {file_path.suffix}")
                return None
            
            return self._load_single_file(file_path)
                
        except Exception as e:
            logger.error(f"Error loading document {file_path}: {e}")
            return None
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if file type is supported"""
        supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.xlsx', '.xls', 
                              '.json', '.csv', '.md', '.markdown'}
        return file_path.suffix.lower() in supported_extensions
    
    def _load_single_file(self, file_path: Path) -> Optional[LangchainDocument]:
        """Load a single file and return as LangchainDocument"""
        
        # Skip if already processed
        file_key = str(file_path.absolute())
        if file_key in self.processed_files:
            logger.debug(f"Skipping already processed file: {file_path.name}")
            return None
        
        try:
            content = self._extract_content(file_path)
            
            if not content or len(content.strip()) == 0:
                logger.warning(f"No content extracted from: {file_path.name}")
                return None
            
            # Create metadata
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "file_type": file_path.suffix.lower().lstrip('.'),
                "file_size": file_path.stat().st_size,
                "processed_at": str(pd.Timestamp.now())
            }
            
            # Add file-specific metadata
            if file_path.suffix.lower() == '.pdf':
                metadata.update(self._get_pdf_metadata(file_path))
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                metadata.update(self._get_excel_metadata(file_path))
            
            # Create LangchainDocument
            document = LangchainDocument(
                page_content=content,
                metadata=metadata
            )
            
            # Mark as processed
            self.processed_files.add(file_key)
            
            return document
            
        except Exception as e:
            logger.error(f"Error processing file {file_path.name}: {e}")
            return None
    
    def _extract_content(self, file_path: Path) -> str:
        """Extract content from file based on its type"""
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_docx(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._extract_excel(file_path)
            elif extension == '.txt':
                return self._extract_text(file_path)
            elif extension == '.json':
                return self._extract_json(file_path)
            elif extension == '.csv':
                return self._extract_csv(file_path)
            elif extension in ['.md', '.markdown']:
                return self._extract_markdown(file_path)
            else:
                logger.warning(f"Unsupported file type: {extension}")
                return ""
                
        except Exception as e:
            logger.error(f"Content extraction failed for {file_path.name}: {e}")
            return ""
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1} from {file_path.name}: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"PDF extraction error for {file_path.name}: {e}")
            
        return text.strip()
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
                        
        except Exception as e:
            logger.error(f"DOCX extraction error for {file_path.name}: {e}")
            
        return text.strip()
    
    def _extract_excel(self, file_path: Path) -> str:
        """Extract text from Excel file"""
        text = ""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if not df.empty:
                        text += f"\n=== Sheet: {sheet_name} ===\n"
                        text += f"Columns: {', '.join(df.columns)}\n"
                        text += f"Rows: {len(df)}\n\n"
                        
                        # Convert to string representation
                        text += df.to_string(index=False, max_rows=100) + "\n"
                        
                        if len(df) > 100:
                            text += f"\n... ({len(df) - 100} more rows)\n"
                            
                except Exception as e:
                    logger.warning(f"Error processing sheet {sheet_name}: {e}")
                    continue
                    
        except Exception as e:
            logger.error(f"Excel extraction error for {file_path.name}: {e}")
            
        return text.strip()
    
    def _extract_text(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
                    
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                return file.read()
                
        except Exception as e:
            logger.error(f"Text extraction error for {file_path.name}: {e}")
            return ""
    
    def _extract_json(self, file_path: Path) -> str:
        """Extract text from JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
                
            # Format JSON nicely
            formatted_json = json.dumps(data, indent=2, ensure_ascii=False)
            
            # Add summary if it's a large JSON
            if len(formatted_json) > 10000:
                summary = f"Large JSON file with {len(str(data))} characters\n"
                summary += f"Top-level keys: {list(data.keys()) if isinstance(data, dict) else 'Not a dict'}\n\n"
                return summary + formatted_json[:5000] + "\n... (truncated)"
            
            return formatted_json
            
        except Exception as e:
            logger.error(f"JSON extraction error for {file_path.name}: {e}")
            return ""
    
    def _extract_csv(self, file_path: Path) -> str:
        """Extract text from CSV file"""
        try:
            # Read CSV with different encodings and separators
            df = None
            
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                for sep in [',', ';', '\t']:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding, sep=sep)
                        if len(df.columns) > 1:  # Successfully parsed
                            break
                    except Exception:
                        continue
                if df is not None and len(df.columns) > 1:
                    break
            
            if df is None or df.empty:
                return ""
            
            # Format CSV content
            text = f"CSV Data from {file_path.name}\n"
            text += f"Columns: {', '.join(df.columns)}\n"
            text += f"Rows: {len(df)}\n\n"
            
            # Add sample data
            sample_size = min(50, len(df))
            text += df.head(sample_size).to_string(index=False)
            
            if len(df) > sample_size:
                text += f"\n... ({len(df) - sample_size} more rows)"
            
            return text
            
        except Exception as e:
            logger.error(f"CSV extraction error for {file_path.name}: {e}")
            return ""
    
    def _extract_markdown(self, file_path: Path) -> str:
        """Extract text from Markdown file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Add file info
            return f"Markdown document: {file_path.name}\n\n{content}"
            
        except Exception as e:
            logger.error(f"Markdown extraction error for {file_path.name}: {e}")
            return ""
    
    def _get_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Get PDF-specific metadata"""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                metadata['pages'] = len(reader.pages)
                
                # Get document info if available
                if reader.metadata:
                    info = reader.metadata
                    metadata['title'] = str(info.get('/Title', ''))
                    metadata['author'] = str(info.get('/Author', ''))
                    metadata['subject'] = str(info.get('/Subject', ''))
                    
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata for {file_path.name}: {e}")
            
        return metadata
    
    def _get_excel_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Get Excel-specific metadata"""
        metadata = {}
        try:
            excel_file = pd.ExcelFile(file_path)
            metadata['sheets'] = excel_file.sheet_names
            metadata['sheet_count'] = len(excel_file.sheet_names)
            
        except Exception as e:
            logger.warning(f"Could not extract Excel metadata for {file_path.name}: {e}")
            
        return metadata
    
    def split_documents(self, documents: List[LangchainDocument]) -> List[LangchainDocument]:
        """Split documents into smaller chunks"""
        if not documents:
            logger.warning("No documents to split")
            return []
        
        logger.info(f"Splitting {len(documents)} documents into chunks...")
        
        try:
            chunks = self.text_splitter.split_documents(documents)
            logger.info(f"Created {len(chunks)} chunks from {len(documents)} documents")
            
            # Add chunk metadata
            for i, chunk in enumerate(chunks):
                chunk.metadata['chunk_id'] = i
                chunk.metadata['chunk_size'] = len(chunk.page_content)
                
                # Ensure required metadata exists
                if 'filename' not in chunk.metadata:
                    chunk.metadata['filename'] = f"chunk_{i}"
                if 'file_type' not in chunk.metadata:
                    chunk.metadata['file_type'] = 'unknown'
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error splitting documents: {e}")
            return []
    
    def get_document_stats(self, documents: List[LangchainDocument]) -> Dict[str, Any]:
        """Get statistics about loaded documents"""
        if not documents:
            return {"total_documents": 0}
        
        stats = {
            'total_documents': len(documents),
            'file_types': {},
            'total_content_length': 0,
            'average_content_length': 0,
            'largest_document': 0,
            'smallest_document': float('inf')
        }
        
        for doc in documents:
            # Count file types
            file_type = doc.metadata.get('file_type', 'unknown')
            stats['file_types'][file_type] = stats['file_types'].get(file_type, 0) + 1
            
            # Calculate content length statistics
            content_length = len(doc.page_content)
            stats['total_content_length'] += content_length
            stats['largest_document'] = max(stats['largest_document'], content_length)
            stats['smallest_document'] = min(stats['smallest_document'], content_length)
        
        if stats['total_documents'] > 0:
            stats['average_content_length'] = stats['total_content_length'] / stats['total_documents']
        
        if stats['smallest_document'] == float('inf'):
            stats['smallest_document'] = 0
        
        return stats
    
    def clear_processed_files(self):
        """Clear the set of processed files to allow reprocessing"""
        self.processed_files.clear()
        logger.info("Cleared processed files cache")
    
    def update_chunk_settings(self, chunk_size: int, chunk_overlap: int):
        """Update chunking parameters"""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        logger.info(f"Updated chunk settings: size={chunk_size}, overlap={chunk_overlap}")
    
    def validate_file_before_processing(self, file_path: Path) -> Dict[str, Any]:
        """Validate a file before processing and return detailed info"""
        validation = {
            "valid": False,
            "reason": "",
            "file_info": {},
            "estimated_processing_time": 0
        }
        
        try:
            if not file_path.exists():
                validation["reason"] = "File does not exist"
                return validation
            
            if not file_path.is_file():
                validation["reason"] = "Path is not a file"
                return validation
            
            if not self._is_supported_file(file_path):
                validation["reason"] = f"Unsupported file type: {file_path.suffix}"
                return validation
            
            # Get file info
            stat = file_path.stat()
            file_size = stat.st_size
            
            validation["file_info"] = {
                "name": file_path.name,
                "size_bytes": file_size,
                "size_mb": file_size / (1024 * 1024),
                "extension": file_path.suffix.lower(),
                "modified": stat.st_mtime
            }
            
            # Check file size limits
            max_size_mb = 100  # 100MB limit
            if file_size > max_size_mb * 1024 * 1024:
                validation["reason"] = f"File too large: {file_size / (1024*1024):.1f}MB (max: {max_size_mb}MB)"
                return validation
            
            # Estimate processing time based on file size and type
            if file_path.suffix.lower() == '.pdf':
                validation["estimated_processing_time"] = file_size / (1024 * 1024) * 2  # 2 seconds per MB for PDF
            elif file_path.suffix.lower() in ['.xlsx', '.xls']:
                validation["estimated_processing_time"] = file_size / (1024 * 1024) * 1.5  # 1.5 seconds per MB for Excel
            else:
                validation["estimated_processing_time"] = file_size / (1024 * 1024) * 0.5  # 0.5 seconds per MB for text
            
            validation["valid"] = True
            validation["reason"] = "File is valid for processing"
            
        except Exception as e:
            validation["reason"] = f"Validation error: {str(e)}"
        
        return validation
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get information about supported file formats"""
        return {
            "document_formats": [".pdf", ".docx", ".doc", ".txt", ".md", ".markdown"],
            "spreadsheet_formats": [".xlsx", ".xls", ".csv"],
            "data_formats": [".json"],
            "descriptions": {
                ".pdf": "Portable Document Format - extracts text and metadata",
                ".docx": "Microsoft Word Document - extracts text, tables, and formatting",
                ".doc": "Legacy Microsoft Word Document",
                ".txt": "Plain Text File - supports multiple encodings",
                ".md": "Markdown Document - preserves formatting",
                ".markdown": "Markdown Document - preserves formatting", 
                ".xlsx": "Microsoft Excel Workbook - extracts all sheets and data",
                ".xls": "Legacy Microsoft Excel Workbook",
                ".csv": "Comma-Separated Values - auto-detects delimiters",
                ".json": "JavaScript Object Notation - preserves structure"
            }
        }
    
    def health_check(self) -> Dict[str, Any]:
        """Perform a health check of the document processor"""
        health = {
            "status": "unknown",
            "text_splitter": False,
            "supported_formats": 0,
            "can_process": False,
            "processed_files_count": len(self.processed_files)
        }
        
        try:
            # Check text splitter
            if self.text_splitter:
                test_doc = LangchainDocument(
                    page_content="This is a test document for health check.",
                    metadata={"test": True}
                )
                chunks = self.text_splitter.split_documents([test_doc])
                health["text_splitter"] = len(chunks) > 0
            
            # Count supported formats
            formats = self.get_supported_formats()
            total_formats = sum(len(fmt_list) for fmt_list in formats.values() if isinstance(fmt_list, list))
            health["supported_formats"] = total_formats
            
            # Test processing capability
            health["can_process"] = health["text_splitter"] and health["supported_formats"] > 0
            
            # Overall status
            if health["can_process"]:
                health["status"] = "healthy"
            else:
                health["status"] = "unhealthy"
                
        except Exception as e:
            health["status"] = "error"
            health["error"] = str(e)
        
        return health